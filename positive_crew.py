import io
import os
import PyPDF2
import re
import streamlit as st

from docx import Document
from docxtpl import DocxTemplate
from openai import OpenAI

from langchain_community.document_loaders import UnstructuredFileLoader

from myfunc.mojafunkcija import positive_login

client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))


# all
def extract_narucilac(text):
    """
    Extracts and returns text found between two markers.
    
    Args:
        text (str): The full text to search within.
        start_marker (str): The text immediately preceding the desired content.
        end_marker (str): The text immediately following the desired content.
    
    Returns:
        str: The extracted text found between start_marker and end_marker.
    """
    pattern = re.compile(rf'{re.escape("Naručilac:")}(.*?){re.escape("Ponuđač:")}', re.S)
    matches = pattern.search(text)
    if matches:
        return matches.group(1).strip()
    else:
        return "No information found between the specified markers."


# all
def extract_datum(text):
    """
    Extracts the first line after a specific block of text using regular expressions
    to allow for flexibility in matching.
    
    Args:
        text (str): The full text to search within.
    
    Returns:
        str: The extracted date or an explicit message indicating the outcome.
    """
    # Regex pattern to match the block ending with "Novi Sad" followed by any amount of whitespace
    # and then capturing the next line of text.
    pattern = re.compile(r"Ponuđač:\s+Positive d\.o\.o\.\s*\n\s*Danila Kiša 5\s*\n\s*Novi Sad\s*\n(.*?)\n", re.DOTALL)
    
    # Search for the pattern in the text
    match = pattern.search(text)
    
    if match:
        # Extract and return the captured group, which should be the date line
        date = match.group(1).strip().replace("godine", "").strip()
        return date
    else:
        return "Date not found after the specified block."


# all
def extract_broj_ponude(text):
    """
    Extracts the number formatted line that appears after "Broj:".
    
    Args:
        text (str): The full text to search within.
    
    Returns:
        str: The extracted number or an explicit message indicating the outcome.
    """
    # Split the text into lines
    lines = text.split('\n')
    
    # Iterate through each line
    for line in lines:
        # Check if the line starts with "Broj: "
        if line.strip().startswith("Broj:"):
            # Extract the part after "Broj: "
            number = line.strip()[6:].strip()  # Remove "Broj: " and any leading/trailing whitespace
            return number
    
    return "Number not found after 'Broj:'"
    

def extract_tip_podrske(text):
    pattern = r"Trenutna telefonska podrška – (\d{1,2}/\d{1,2})"
    match = re.search(pattern, text)

    if match:
        extracted_hours = match.group(1)  # Group 1 contains the matched digits pattern
        return extracted_hours
    else:
        print("No match found - kopiraj >> u svemu prema specifikaciji navedenoj u Ponudi Davaoca usluga broj ___________, izdatoj _________ godine >> u tekst ugovora")


# all
def extract_broj_meseci(text):
    """
    Extracts the number of months from a sentence indicating the contract period.
    
    Args:
        text (str): The full text to search within.
    
    Returns:
        str: The extracted number of months or an explicit message indicating the outcome.
    """
    # Regex pattern to match the sentence and capture the number of months
    pattern1 = re.compile(r"Ugovor se sklapa na period od (\d+) meseci")
    pattern2 = re.compile(r"Ugovor se sklapa na (\d+) meseci")

    # Search for the pattern in the text
    match1 = pattern1.search(text)
    match2 = pattern2.search(text)
    
    if match1:
        # Extract and return the number of months
        return match1.group(1)
    elif match2:
        return match2.group(1)
    else:
        return "Number of months not found."


# ESS, PCRM
def extract_usluga(text):
    """
    Extracts a block of text that details services offered, located between
    "Usluga podrazumeva:" and "Cena usluge na mesečnom nivou".
    
    Args:
        text (str): The full text to search within.
    
    Returns:
        str: The extracted block of text or a message indicating it was not found.
    """
    # Define a regex pattern to capture everything between the two markers


    # pattern = re.compile(r"Usluga podrazumeva:(.*?)Cena usluge na mesečnom nivou", re.S)
    pattern = re.compile(r"Usluga PC Radno Mesto uključuje:(.*?)Cena usluge na mesečnom nivou", re.S)
    
    # Search for the pattern in the text
    match = pattern.search(text)
    
    if match:
        # Return the captured block of text, stripping leading/trailing whitespace
        return match.group(1).strip()
    else:
        return "Service details not found between the specified markers."


# ESS, PCRM
def extract_detalji_usluge(text):
    """
    Extracts specific details from the service description block.
    
    Args:
        text (str): The block of text detailing the services.
    
    Returns:
        tuple: Contains the extracted "Sophos A or E", "Server" bool, and "Phone Availability".
    """
    # 1. Extract "Advanced" or "Essentials"
    sophos_a_or_e_match = re.search(r"Sophos Central Intercept X (Advanced|Essentials)", text)
    sophos_a_or_e = sophos_a_or_e_match.group(1) if sophos_a_or_e_match else "Not Found"
    
    # 2. Check for the word "Server"
    server = bool(re.search(r"\bServer\b", text))
    
    # 3. Extract "8/5" or similar pattern
    phone_availability_match = re.search(r"Trenutna telefonska podrška – (\d+/\d+)", text)
    phone_availability = phone_availability_match.group(1) if phone_availability_match else "Not Found"
    
    return sophos_a_or_e, server, phone_availability


# all
def extract_direktor(text):
    """
    Extracts the director's name placeholder or text between "S poštovanjem," and "Ponudu prihvata",
    handling variable whitespace and line breaks.
    
    Args:
        text (str): The full text to search within.
    
    Returns:
        str: The extracted text or a message indicating it was not found.
    """
    # Define a regex pattern to match the area of interest with flexible whitespace
    pattern = re.compile(r"S poštovanjem,\s*(___.*?___)\s*\n\s*Ponudu prihvata", re.S)
    
    # Search for the pattern in the text
    match = pattern.search(text)
    
    if match:
        # Return the captured text
        return match.group(1)
    else:
        return "Director's name or placeholder not found."


# PCRM
def extract_cena_eur(text):
    """
    Extracts the number that appears after "odnosno" and immediately before "eur" in the text.
    
    Args:
        text (str): The full text to search within.
    
    Returns:
        str: The extracted number or a message indicating it was not found.
    """
    # Define a regex pattern to capture a number that follows "odnosno" and precedes "eur"
    pattern = re.compile(r"odnosno\s+(\d+)\s*eur", re.IGNORECASE)
    
    # Search for the pattern in the text
    match = pattern.search(text)
    
    if match:
        # Return the captured number
        return match.group(1)
    else:
        return "Specific number before 'eur' not found."



# ORM
def extract_MS_365(text):
    """
    Extracts the text between "rešenjem mesečno po korisniku." and "Cena usluge na mesečnom nivou".
    
    Args:
        text (str): The full text to search within.
    
    Returns:
        str: The extracted text or a message indicating it was not found.
    """
    # Define a regex pattern to capture everything between the specified phrases
    pattern1 = re.compile(r"rešenjem mesečno po korisniku\.(.*?)Cena usluge na mesečnom nivou", re.S)
    pattern2 = re.compile(r"Usluga podrazumeva\:(.*?)Cena usluge na mesečnom nivou", re.S)
    pattern3 = re.compile(r"Komponente usluge\:(.*?)Cena usluge na mesečnom nivou", re.S)

    # Search for the pattern in the text
    match1 = pattern1.search(text)
    match2 = pattern2.search(text)
    match3 = pattern3.search(text)
    if match1:
        # Return the captured block of text, stripping leading/trailing whitespace
        return match1.group(1).strip()
    elif match2:
        return match2.group(1).strip()
    elif match3:
        return match3.group(1).strip()
    else:
        return "Text not found between the specified phrases."



# SEL
def extract_produkt_i_cena(text):
    """
    Extracts specific details from a structured product list within the text.
    
    Args:
        text (str): The full text to search within.
    
    Returns:
        tuple: Contains the extracted product description and the price per unit.
    """
    # Step 1: Extract the relevant block of text
    pattern_block = re.compile(r"R\.br\.\s+Opis\s+J\.m\.\s+Kol\.\s+Cena\s+Suma\s+PDV\s+Ukupno(.*?)Ukupno za uplatu", re.S)
    block_match = pattern_block.search(text)
    
    if not block_match:
        return ("Text block not found", "")

    text_block = block_match.group(1).strip()

    # Step 2: Extract the product description
    # pattern_desc = re.compile(r"(\D+?)\s+kom")

    # MS LIC
    pattern_desc = re.compile(r"\d+\s+(.*?)\s+kom")
    desc_match = pattern_desc.search(text_block)
    product_description = desc_match.group(1).strip() if desc_match else "Product description not found"

    # Step 3: Extract the price per unit
    pattern_price = re.compile(r"kom\s+\d+\s+([\d\.]+)")
    price_match = pattern_price.search(text_block)
    price_per_unit = price_match.group(1).strip() if price_match else "Price per unit not found"

    return (product_description, price_per_unit)



def format_address(multi_line_address, director_name):
    # Split the address into lines and strip any extra whitespace
    lines = [line.strip() for line in multi_line_address.split('\n') if line.strip()]
    
    # Extract the relevant parts
    company_name = lines[0]
    address = lines[1]
    city = lines[2]
    mb = lines[3].split(': ')[1]  # Assuming 'MB:' is followed by the number
    pib = lines[4].split(': ')[1]  # Assuming 'PIB:' is followed by the number
    
    # Format into the desired single line
    formatted_address = f"{company_name}, {city}, ul. {address}, MB: {mb}, PIB: {pib}, koga zastupa direktor/ovlašćeno lice {director_name}"
    
    return formatted_address


def replace_content_ORM():
    multi_line_address = extract_narucilac(st.session_state.uploaded_file_content)
    ime_firme = multi_line_address.splitlines()[0]
    director_name = extract_direktor(st.session_state.uploaded_file_content)

    # Format the address
    formatted_address = format_address(multi_line_address, director_name)

    docx_path = r'C:\Users\nemanja.perunicic\OneDrive - Positive doo\Desktop\allIn1\crewai\ugovor ORM.docx'
    old_text_start = "POSLOVNO IME"  # A unique part of the paragraph to be replaced

    replacements = [extract_broj_ponude(st.session_state.uploaded_file_content), extract_datum(st.session_state.uploaded_file_content)]

    broj_meseci = extract_broj_meseci(st.session_state.uploaded_file_content)

    doc = Document(docx_path)
    for paragraph in doc.paragraphs:
        if old_text_start in paragraph.text:
            # Clear existing paragraph
            for run in paragraph.runs:
                run.clear()
            # Insert new text
            new_run = paragraph.add_run()  # Create a new run
            # Add the bold part
            new_run.text = ime_firme
            new_run.bold = True
            # Add the rest of the address in a separate run
            rest_of_address = formatted_address[len(ime_firme):]
            paragraph.add_run(rest_of_address)
            break
    
    placeholder_pattern = re.compile(r"_{8,}")  # Matches 8 or more underscores

    for paragraph in doc.paragraphs:
        placeholder_pattern = re.compile(r"_{8,}")  # Matches 8 or more underscores
        if "specifikaciji navedenoj u Ponudi Davaoca usluga" in paragraph.text:
            # Replace the placeholders sequentially
            original_text = paragraph.text
            for replacement in replacements:
                original_text = placeholder_pattern.sub(replacement, original_text, 1)  # Replace one at a time
            paragraph.text = original_text
            break

    for paragraph in doc.paragraphs:
        meseci_pattern = re.compile(r"\d+\s*meseci")
        if meseci_pattern.search(paragraph.text):
            updated_text = meseci_pattern.sub(f"{broj_meseci} meseci", paragraph.text)
            paragraph.text = updated_text

    cena_eur = float(extract_cena_eur(st.session_state.uploaded_file_content))
    broj_meseci = int(broj_meseci)
    count = 0
    for paragraph in doc.paragraphs:
        if "0,00 €" in paragraph.text:
            if count == 0:
                paragraph.text = paragraph.text.replace("0,00", str(cena_eur * broj_meseci), 1)
                paragraph.text = paragraph.text.replace("0,00", str(cena_eur * broj_meseci * 1.2), 1)

                count += 1
            if count == 1:
                paragraph.text = paragraph.text.replace("0,00", str(cena_eur), 1)
                paragraph.text = paragraph.text.replace("0,00", str(cena_eur * 1.2), 1)


    for paragraph in doc.paragraphs:
        if str(cena_eur * broj_meseci * 1.2) in paragraph.text:
            paragraph.text = re.sub(r'\s+', ' ', paragraph.text)
    doc.save(docx_path.replace('.docx', '_modified.docx'))




def replace_content_PCRM():
    multi_line_address = extract_narucilac(st.session_state.uploaded_file_content)
    ime_firme = multi_line_address.splitlines()[0]
    director_name = extract_direktor(st.session_state.uploaded_file_content)

    # Format the address
    formatted_address = format_address(multi_line_address, director_name)
    docx_path = r'C:\Users\nemanja.perunicic\OneDrive - Positive doo\Desktop\allIn1\crewai\ugovor PCRM.docx'
    old_text_start = "POSLOVNO IME"  # A unique part of the paragraph to be replaced

    replacements = [extract_broj_ponude(st.session_state.uploaded_file_content), extract_datum(st.session_state.uploaded_file_content)]

    broj_meseci = extract_broj_meseci(st.session_state.uploaded_file_content)

    doc = Document(docx_path)
    for paragraph in doc.paragraphs:
        if old_text_start in paragraph.text:
            # Clear existing paragraph
            for run in paragraph.runs:
                run.clear()
            # Insert new text
            new_run = paragraph.add_run()  # Create a new run
            # Add the bold part
            new_run.text = ime_firme
            new_run.bold = True
            # Add the rest of the address in a separate run
            rest_of_address = formatted_address[len(ime_firme):]
            paragraph.add_run(rest_of_address)
            break
    
    placeholder_pattern = re.compile(r"_{8,}")  # Matches 8 or more underscores

    for paragraph in doc.paragraphs:
        placeholder_pattern = re.compile(r"_{8,}")  # Matches 8 or more underscores
        if "specifikaciji navedenoj u Ponudi Davaoca usluga" in paragraph.text:
            # Replace the placeholders sequentially
            original_text = paragraph.text
            for replacement in replacements:
                original_text = placeholder_pattern.sub(replacement, original_text, 1)  # Replace one at a time
            paragraph.text = original_text
            break


    tip_podrske = extract_tip_podrske(st.session_state.uploaded_file_content)
    # Assuming 'doc.paragraphs' and 'broj_meseci' are defined
    for paragraph in doc.paragraphs:
        # Check if the paragraph contains the simplified identification text
        if "Pick 24 ili 36" in paragraph.text:
            # Determine the replacement based on 'broj_meseci'
            godina_replacement = "prve" if broj_meseci == "24" else "druge"
            # Replace "prve/druge" with the determined replacement
            paragraph.text = re.sub(r"\bprve/druge\b", godina_replacement, paragraph.text)
            # Remove the simplified identification text
            paragraph.text = re.sub(r"Pick 24 ili 36", "", paragraph.text).strip()

    # Ensure to recheck the paragraph's text to confirm the changes are applied as expected.
        if "Pick 24/7 ili 8/5" in paragraph.text:
            # Determine the replacement based on 'broj_meseci'
            dan_replacement = "jednog radnog" if tip_podrske == "24/7" else "dva radna"
            # Replace "prve/druge" with the determined replacement
            paragraph.text = re.sub(r"\bjednog radnog /dva radna\b", dan_replacement, paragraph.text)
            # Remove the simplified identification text
            paragraph.text = re.sub(r"Pick 24/7 ili 8/5", "", paragraph.text).strip()

        if "Obezbeđivanje podrške – 24/7 ili 8/5" in paragraph.text:
            paragraph.text = re.sub(r"\b24/7 ili 8/5\b", tip_podrske, paragraph.text)

        if "Picks 24/7" in paragraph.text:
            if tip_podrske == "24/7":
                paragraph.text = re.sub(r"Picks 24/7", "", paragraph.text).strip()
            else:
                paragraph.text = ""
        if "Picks 8/5" in paragraph.text:
            if tip_podrske == "8/5":
                paragraph.text = re.sub(r"Picks 8/5", "", paragraph.text).strip()
            else:
                paragraph.text = ""

    # Assuming 'doc.paragraphs' and 'broj_meseci' are defined earlier in your code
    for paragraph in doc.paragraphs:
        # This pattern now accounts for both Cyrillic 'а' and Latin 'a', and optional 'i' following "meseca"
        inclusive_pattern = re.compile(r"12/24/36\s*mesec[аa]/?i?")
        if inclusive_pattern.search(paragraph.text):
            # Replace with 'broj_meseci meseci', assuming 'broj_meseci' is the desired replacement
            updated_text = inclusive_pattern.sub(f"{broj_meseci} meseci", paragraph.text)
            paragraph.text = updated_text

    cena_eur = float(extract_cena_eur(st.session_state.uploaded_file_content))
    broj_meseci = int(broj_meseci)
    count = 0
    for paragraph in doc.paragraphs:
        if "0,00 €" in paragraph.text:
            if count == 0:
                paragraph.text = paragraph.text.replace("0,00", str(cena_eur * broj_meseci), 1)
                paragraph.text = paragraph.text.replace("0,00", str(cena_eur * broj_meseci * 1.2), 1)

                count += 1
            if count == 1:
                paragraph.text = paragraph.text.replace("0,00", str(cena_eur), 1)
                paragraph.text = paragraph.text.replace("0,00", str(cena_eur * 1.2), 1)


    for paragraph in doc.paragraphs:
        if str(cena_eur * broj_meseci * 1.2) in paragraph.text:
            paragraph.text = re.sub(r'\s+', ' ', paragraph.text)
    doc.save(docx_path.replace('.docx', '_modified.docx'))




def replace_content_ESS(string_without_empty_lines):
    multi_line_address = extract_narucilac(st.session_state.uploaded_file_content)
    ime_firme = multi_line_address.splitlines()[0]
    director_name = extract_direktor(st.session_state.uploaded_file_content)

    # Format the address
    formatted_address = format_address(multi_line_address, director_name)

    docx_path = r'C:\Users\nemanja.perunicic\OneDrive - Positive doo\Desktop\allIn1\crewai\ugovor ESS.docx'
    old_text_start = "POSLOVNO IME"  # A unique part of the paragraph to be replaced

    replacements = [extract_broj_ponude(st.session_state.uploaded_file_content), extract_datum(st.session_state.uploaded_file_content)]

    broj_meseci = extract_broj_meseci(st.session_state.uploaded_file_content)

    doc = Document(docx_path)
    for paragraph in doc.paragraphs:
        if old_text_start in paragraph.text:
            # Clear existing paragraph
            for run in paragraph.runs:
                run.clear()
            # Insert new text
            new_run = paragraph.add_run()  # Create a new run
            # Add the bold part
            new_run.text = ime_firme
            new_run.bold = True
            # Add the rest of the address in a separate run
            rest_of_address = formatted_address[len(ime_firme):]
            paragraph.add_run(rest_of_address)
            break
    
    placeholder_pattern = re.compile(r"_{8,}")  # Matches 8 or more underscores

    for paragraph in doc.paragraphs:
        placeholder_pattern = re.compile(r"_{8,}")  # Matches 8 or more underscores
        if "specifikaciji navedenoj u Ponudi Davaoca usluga" in paragraph.text:
            # Replace the placeholders sequentially
            original_text = paragraph.text
            for replacement in replacements:
                original_text = placeholder_pattern.sub(replacement, original_text, 1)  # Replace one at a time
            paragraph.text = original_text
            break

    for paragraph in doc.paragraphs:
        meseci_pattern = re.compile(r"\d+\s*meseci")
        if meseci_pattern.search(paragraph.text):
            updated_text = meseci_pattern.sub(f"{broj_meseci} meseci", paragraph.text)
            paragraph.text = updated_text

    string_without_empty_lines = string_without_empty_lines.split("\n")
    if "Essentials" in string_without_empty_lines[0]:
        for paragraph in doc.paragraphs:
            if "5,00 €" and "6,50 €" in paragraph.text:
                remove_paragraph(paragraph)
    elif "Advanced" in string_without_empty_lines[0]:
        for paragraph in doc.paragraphs:
            if "2,50 €" and "3,00 €" in paragraph.text:
                remove_paragraph(paragraph)
    if "Server" not in string_without_empty_lines[1]:
        for paragraph in doc.paragraphs:
            if "8,50 €" and "10,20 €" in paragraph.text:
                remove_paragraph(paragraph)
            elif "Obim pružanja usluge zavisi od broja klijentskih korisnika i servera" in paragraph.text:
                paragraph.text = re.sub(r"i servera", "", paragraph.text).strip()
    else:
        for paragraph in doc.paragraphs:
            if "8,50 €" and "10,20 €" in paragraph.text:
                paragraph.text = re.sub(paragraph.text, paragraph.text, paragraph.text).strip()
    doc.save(docx_path.replace('.docx', '_modified.docx'))


def replace_content_SMS(string_without_empty_lines):
    multi_line_address = extract_narucilac(st.session_state.uploaded_file_content)
    ime_firme = multi_line_address.splitlines()[0]
    director_name = extract_direktor(st.session_state.uploaded_file_content)

    # Format the address
    formatted_address = format_address(multi_line_address, director_name)

    docx_path = r'C:\Users\nemanja.perunicic\OneDrive - Positive doo\Desktop\allIn1\crewai\ugovor SMS.docx'
    old_text_start = "POSLOVNO IME"  # A unique part of the paragraph to be replaced

    replacements = [extract_broj_ponude(st.session_state.uploaded_file_content), extract_datum(st.session_state.uploaded_file_content)]

    broj_meseci = extract_broj_meseci(st.session_state.uploaded_file_content)

    doc = Document(docx_path)
    for paragraph in doc.paragraphs:
        if old_text_start in paragraph.text:
            # Clear existing paragraph
            for run in paragraph.runs:
                run.clear()
            # Insert new text
            new_run = paragraph.add_run()  # Create a new run
            # Add the bold part
            new_run.text = ime_firme
            new_run.bold = True
            # Add the rest of the address in a separate run
            rest_of_address = formatted_address[len(ime_firme):]
            paragraph.add_run(rest_of_address)
            break
    
    placeholder_pattern = re.compile(r"_{8,}")  # Matches 8 or more underscores
    for paragraph in doc.paragraphs:
        placeholder_pattern = re.compile(r"_{8,}")  # Matches 8 or more underscores
        if "specifikaciji navedenoj u Ponudi Davaoca usluga" in paragraph.text:
            # Replace the placeholders sequentially
            original_text = paragraph.text
            for replacement in replacements:
                original_text = placeholder_pattern.sub(replacement, original_text, 1)  # Replace one at a time
            paragraph.text = original_text
            break

    for paragraph in doc.paragraphs:
        meseci_pattern = re.compile(r"\d+\s*meseci")
        if meseci_pattern.search(paragraph.text):
            updated_text = meseci_pattern.sub(f"{broj_meseci} meseci", paragraph.text)
            paragraph.text = updated_text

    string_without_empty_lines = string_without_empty_lines.split("\n")
    if "Essentials" in string_without_empty_lines[0]:
        for paragraph in doc.paragraphs:
            if "5,00 €" and "6,50 €" in paragraph.text:
                remove_paragraph(paragraph)
    elif "Advanced" in string_without_empty_lines[0]:
        for paragraph in doc.paragraphs:
            if "2,50 €" and "3,00 €" in paragraph.text:
                remove_paragraph(paragraph)
    if "Server" not in string_without_empty_lines[1]:
        for paragraph in doc.paragraphs:
            if "8,50 €" and "10,20 €" in paragraph.text:
                remove_paragraph(paragraph)
            elif "Obim pružanja usluge zavisi od broja klijentskih korisnika i servera" in paragraph.text:
                paragraph.text = re.sub(r"i servera", "", paragraph.text).strip()
    else:
        for paragraph in doc.paragraphs:
            if "8,50 €" and "10,20 €" in paragraph.text:
                paragraph.text = re.sub(paragraph.text, paragraph.text, paragraph.text).strip()


    tip_podrske = extract_tip_podrske(st.session_state.uploaded_file_content)
    # Assuming 'doc.paragraphs' and 'broj_meseci' are defined
    for paragraph in doc.paragraphs:
        # Check if the paragraph contains the simplified identification text
        if "Pick 24 ili 36" in paragraph.text:
            # Determine the replacement based on 'broj_meseci'
            godina_replacement = "prve" if broj_meseci == "24" else "druge"
            # Replace "prve/druge" with the determined replacement
            paragraph.text = re.sub(r"\bprve/druge\b", godina_replacement, paragraph.text)
            # Remove the simplified identification text
            paragraph.text = re.sub(r"Pick 24 ili 36", "", paragraph.text).strip()

    # Ensure to recheck the paragraph's text to confirm the changes are applied as expected.
        if "Pick 24/7 ili 8/5" in paragraph.text:
            # Determine the replacement based on 'broj_meseci'
            dan_replacement = "jednog radnog" if tip_podrske == "24/7" else "dva radna"
            # Replace "prve/druge" with the determined replacement
            paragraph.text = re.sub(r"\bjednog radnog /dva radna\b", dan_replacement, paragraph.text)
            # Remove the simplified identification text
            paragraph.text = re.sub(r"Pick 24/7 ili 8/5", "", paragraph.text).strip()

        if "Obezbeđivanje podrške – 24/7 ili 8/5" in paragraph.text:
            paragraph.text = re.sub(r"\b24/7 ili 8/5\b", tip_podrske, paragraph.text)

        if "Picks 24/7" in paragraph.text:
            if tip_podrske == "24/7":
                paragraph.text = re.sub(r"Picks 24/7", "", paragraph.text).strip()
            else:
                paragraph.text = ""
        if "Picks 8/5" in paragraph.text:
            if tip_podrske == "8/5":
                paragraph.text = re.sub(r"Picks 8/5", "", paragraph.text).strip()
            else:
                paragraph.text = ""

    cena_eur = float(extract_cena_eur(st.session_state.uploaded_file_content))
    broj_meseci = int(broj_meseci)
    count = 0
    for paragraph in doc.paragraphs:
        if "0,00 €" in paragraph.text:
            if count == 0:
                paragraph.text = paragraph.text.replace("0,00", str(cena_eur * broj_meseci), 1)
                paragraph.text = paragraph.text.replace("0,00", str(cena_eur * broj_meseci * 1.2), 1)

                count += 1
            if count == 1:
                paragraph.text = paragraph.text.replace("0,00", str(cena_eur), 1)
                paragraph.text = paragraph.text.replace("0,00", str(cena_eur * 1.2), 1)            
    doc.save(docx_path.replace('.docx', '_modified.docx'))



def replace_content_SEL():
    multi_line_address = extract_narucilac(st.session_state.uploaded_file_content)
    ime_firme = multi_line_address.splitlines()[0]
    director_name = extract_direktor(st.session_state.uploaded_file_content)

    # Format the address
    formatted_address = format_address(multi_line_address, director_name)

    docx_path = r'C:\Users\nemanja.perunicic\OneDrive - Positive doo\Desktop\allIn1\crewai\ugovor SEL.docx'
    old_text_start = "POSLOVNO IME"  # A unique part of the paragraph to be replaced

    replacements = [extract_broj_ponude(st.session_state.uploaded_file_content), extract_datum(st.session_state.uploaded_file_content)]

    broj_meseci = extract_broj_meseci(st.session_state.uploaded_file_content)

    doc = Document(docx_path)
    for paragraph in doc.paragraphs:
        if old_text_start in paragraph.text:
            # Clear existing paragraph
            for run in paragraph.runs:
                run.clear()
            # Insert new text
            new_run = paragraph.add_run()  # Create a new run
            # Add the bold part
            new_run.text = ime_firme
            new_run.bold = True
            # Add the rest of the address in a separate run
            rest_of_address = formatted_address[len(ime_firme):]
            paragraph.add_run(rest_of_address)
            break
    
    placeholder_pattern = re.compile(r"_{8,}")  # Matches 8 or more underscores

    for paragraph in doc.paragraphs:
        placeholder_pattern = re.compile(r"_{8,}")  # Matches 8 or more underscores
        if "specifikaciji navedenoj u Ponudi Davaoca usluga" in paragraph.text:
            # Replace the placeholders sequentially
            original_text = paragraph.text
            for replacement in replacements:
                original_text = placeholder_pattern.sub(replacement, original_text, 1)  # Replace one at a time
            paragraph.text = original_text
            break

    for paragraph in doc.paragraphs:
        meseci_pattern = re.compile(r"\d+\s*meseci")
        if meseci_pattern.search(paragraph.text):
            updated_text = meseci_pattern.sub(f"{broj_meseci} meseci", paragraph.text)
            paragraph.text = updated_text
    
    produkt, cena = extract_produkt_i_cena(st.session_state.uploaded_file_content)
    all_products = ["Sophos Central Intercept X Essentials", "Server Standard", "Intercept X", "Central Intercept X Advanced for Server"]
   
    counter = 0
    x = False
    for paragraph in doc.paragraphs:
        if "Korisnik vrši plaćanje mesečnog zaduženja u iznosu od:" in paragraph.text:
            x = True
            continue
        if x:
            if produkt != all_products[counter]:
                remove_paragraph(paragraph)
            elif produkt == all_products[counter]:
                print(all_products[counter])
            counter += 1

        if counter == len(all_products):
            break

    doc.save(docx_path.replace('.docx', '_modified.docx'))



def remove_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None


# za bullet points
def generate_document(template_path, output_path, bullet_points):
    tpl = DocxTemplate(template_path)
    context = {
        'bullets': bullet_points
    }
    tpl.render(context)
    tpl.save(output_path)


def main():
    if "uploaded_file_content" not in st.session_state:
        st.session_state.uploaded_file_content = None
        

    st.title("Positive 👯 Offer2Contract")
    with st.sidebar:
        st.markdown("ver 1.0")

        ponuda = st.file_uploader(
            "Uploadujte ponudu",
            key="upload_file",
            type=["txt", "docx", "pdf"])

        if ponuda is not None:
            with io.open(ponuda.name, "wb") as file:
                file.write(ponuda.getbuffer())

            if ".pdf" in ponuda.name:
                pdf_reader = PyPDF2.PdfReader(ponuda)
                num_pages = len(pdf_reader.pages)
                text_content = ""

                for page in range(num_pages):
                    page_obj = pdf_reader.pages[page]
                    text_content += page_obj.extract_text()

                text_content = re.sub(r"(?<=\b\w) (?=\w\b)", "", text_content.replace("•", ""))
                with io.open("temp.txt", "w", encoding="utf-8-sig") as f:
                    f.write(text_content)

                loader = UnstructuredFileLoader("temp.txt", encoding="utf-8-sig")
            else:
                loader = UnstructuredFileLoader(file_path=ponuda.name, encoding="utf-8-sig")

            st.session_state.uploaded_file_content = loader.load()[0].page_content

    

    if st.session_state.uploaded_file_content:
        new_content = extract_MS_365(st.session_state.uploaded_file_content) + " (radnim danima od 8 do 16):"
        lines = new_content.split("\n")
        non_empty_lines = [line for line in lines if line.strip() != ""]

        string_without_empty_lines = "\n".join(non_empty_lines)
        #print(st.session_state.uploaded_file_content)
        if "ORM" in ponuda.name:
            replace_content_ORM()
            docx_path = r'C:\Users\nemanja.perunicic\OneDrive - Positive doo\Desktop\allIn1\crewai\ugovor ORM_modified.docx'
        elif "PCRM" in ponuda.name:
            replace_content_PCRM()
            docx_path = r'C:\Users\nemanja.perunicic\OneDrive - Positive doo\Desktop\allIn1\crewai\ugovor PCRM_modified.docx'
        elif "ESS" in ponuda.name:
            replace_content_ESS(string_without_empty_lines)
            docx_path = r'C:\Users\nemanja.perunicic\OneDrive - Positive doo\Desktop\allIn1\crewai\ugovor ESS_modified.docx'
        elif "SMS" in ponuda.name:
            replace_content_SMS(string_without_empty_lines)
            docx_path = r'C:\Users\nemanja.perunicic\OneDrive - Positive doo\Desktop\allIn1\crewai\ugovor SMS_modified.docx'
        elif "SEL" in ponuda.name:
            replace_content_SEL()
            docx_path = r'C:\Users\nemanja.perunicic\OneDrive - Positive doo\Desktop\allIn1\crewai\ugovor SEL_modified.docx'


        # Example usage
        generate_document(docx_path, docx_path.replace('.docx', '_output.docx'), string_without_empty_lines.split('\n'))
        if "ORM" in ponuda.name:
            os.remove("ugovor ORM_modified.docx")
        elif "PCRM" in ponuda.name:
            os.remove("ugovor PCRM_modified.docx")
        elif "ESS" in ponuda.name:
            os.remove("ugovor ESS_modified.docx")
        elif "SMS" in ponuda.name:
            os.remove("ugovor SMS_modified.docx")
        elif "SEL" in ponuda.name:
            os.remove("ugovor SEL_modified.docx")

        st.info("Ugovor je generisan")
        st.divider()
        with st.expander("Sadržaj ponude (za proveru)"):
            st.write(st.session_state.uploaded_file_content)


deployment_environment = os.environ.get("DEPLOYMENT_ENVIRONMENT")

if deployment_environment == "Streamlit":
    name, authentication_status, username = positive_login(main, " ")
else:
    if __name__ == "__main__":
        main()
